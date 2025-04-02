#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Aug 24 11:03:34 2023

@author: mikewoodward
"""

from docx import Document

document = Document('Demo1.docx')

# Properties
print("Properties")
print("==========")
print("Properties title: "
      f"{document.core_properties.title}")
print("Properties author: "
      f"{document.core_properties.author}")
print("Properties subject: "
      f"{document.core_properties.subject}")

# Get the paragraph information
print("Zeroth paragraph")
print("================")
print("Paragraph 0 style: "
      f"{document.paragraphs[0].style}")
print("Paragraph 0 text: "
      f"{document.paragraphs[0].text}")

# Get table information
print("Zeroth Table")
print("============")
print("Table rows: "
      f"{len(document.tables[0].rows)}")
print("Table columns: "
      f"{len(document.tables[0].columns)}")
print("Cell 1,1: "
      f"{document.tables[0].rows[1].cells[1].text}")
print("Cell 2,2: "
      f"{document.tables[0].columns[2].cells[2].text}")

# Get the footer
print("Footer")
print("======")
print("The footer is: "
      f"{document.sections[0].footer.paragraphs[0].text}")
