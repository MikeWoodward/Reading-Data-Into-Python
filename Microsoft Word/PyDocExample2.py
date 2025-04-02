#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Aug 24 11:03:34 2023

@author: mikewoodward
"""

from docx import Document

filename_root = 'Demo2'
document = Document(filename_root + '.docx')

# Replacing $COMPANY_IDENTIFIER, $DATE_PREPARED,
# and $GENERIC_1
company_name = "Mike Co., Ltd"
date_prepared = "April 1st, 2026"
generic_1 = ("This was prepared unter the terms "
             "of the contract DS457/A.")
for paragraph in document.paragraphs:
    if '$COMPANY_IDENTIFIER' in paragraph.text:
        paragraph.text = (
            paragraph
            .text
            .replace('$COMPANY_IDENTIFIER',
                     company_name))
    if '$DATE_PREPARED' in paragraph.text:
        paragraph.text = (
            paragraph
            .text
            .replace('$DATE_PREPARED',
                     date_prepared))
    if '$GENERIC_1' in paragraph.text:
        paragraph.text = (paragraph
                          .text
                          .replace('$GENERIC_1',
                                   generic_1))

# Replacing table cells
cell_1_1 = "11"
cell_2_2 = "22.0"
for table in document.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if '$CELL_1_1' in paragraph.text:
                    paragraph.text = (
                        paragraph
                        .text
                        .replace("$CELL_1_1",
                                 cell_1_1))
                if '$CELL_2_2' in paragraph.text:
                    paragraph.text = (
                        paragraph
                        .text
                        .replace("$CELL_2_2",
                                 cell_2_2))

# Inserting a paragraph at the end
p = document.add_paragraph(
    "A paragraph at the end of the document ")
p.add_run('some bold text').bold = True
p.add_run(' and ')
p.add_run('some italic text.').italic = True

# Save the document to a new name
document.save(filename_root + '_1' + '.docx')
