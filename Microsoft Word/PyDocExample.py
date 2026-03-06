#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Aug 24 11:03:34 2023

@author: mikewoodward
"""

from docx import Document

document = Document('Demo.docx')

# Properties
print("Properties")
print("==========")
print("Properties title: "
      f"{document.core_properties.title}")
print("Properties author: "
      f"{document.core_properties.author}")
print("Properties subject: "
      f"{document.core_properties.subject}")

# Get the paragraph information
print("Zeroth paragraph")
print("================")
print("Paragraph 0 style: "
      f"{document.paragraphs[0].style}")
print("Paragraph 0 text: "
      f"{document.paragraphs[0].text}")

# Get the table of content information

# Get table information
print("Zeroth Table")
print("============")
print("Table rows: "
      f"{len(document.tables[0].rows)}")
print("Table columns: "
      f"{len(document.tables[0].columns)}")
# Reference the table and section first
tab = document.tables[0]
sec = document.sections[0]

# Extract cell text
c1 = tab.rows[1].cells[1].text
print(f"Cell 1,1: {c1}")

c2 = tab.columns[2].cells[2].text
print(f"Cell 2,2: {c2}")

# Get the footer
print("Footer\n======")
f_text = sec.footer.paragraphs[0].text
print(f"The footer is: {f_text}")
